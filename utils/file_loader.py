import os
import json
from PyPDF2 import PdfReader
from docx import Document

def chunk_text_with_metadata(lines, chunk_size=500, overlap=50, page=1):
    """
    lines: list of (text, page, line_num) tuples
    Returns: list of dicts with chunk text, start page, start line
    """
    chunks = []
    i = 0
    while i < len(lines):
        chunk_lines = lines[i:i + chunk_size]
        # Join with newline to preserve line structure for best line extraction
        chunk_text = "\n".join([l[0] for l in chunk_lines])
        start_page = chunk_lines[0][1] if chunk_lines else page
        start_line = chunk_lines[0][2] if chunk_lines else 1
        chunks.append({
            "text": chunk_text,
            "page": start_page,
            "line_num": start_line
        })
        i += chunk_size - overlap
    return chunks

class FileLoader:
    def __init__(self, folder_path, cache_file="chunks_cache.txt"):
        self.folder_path = folder_path
        self.supported_extensions = ['.txt', '.pdf', '.docx']
        self.cache_file = cache_file

    def _load_txt(self, file_path):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _load_pdf(self, file_path):
        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        return text

    def _load_docx(self, file_path):
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def _process_and_chunk_documents(self):
        """
        Loads and chunks all supported documents in the folder.
        Returns a list of (filename, chunk_text, page, line_num) tuples.
        """
        chunks = []
        for root, _, files in os.walk(self.folder_path):
            for file in files:
                ext = os.path.splitext(file)[-1].lower()
                if ext in self.supported_extensions:
                    file_path = os.path.join(root, file)
                    try:
                        lines = []
                        if ext == '.txt':
                            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                for idx, line in enumerate(f, 1):
                                    lines.append((line.strip(), 1, idx))
                        elif ext == '.pdf':
                            reader = PdfReader(file_path)
                            for pageno, page in enumerate(reader.pages, start=1):
                                text = page.extract_text() or ''
                                for idx, line in enumerate(text.split('\n'), 1):
                                    lines.append((line.strip(), pageno, idx))
                        elif ext == '.docx':
                            doc = Document(file_path)
                            for idx, para in enumerate(doc.paragraphs, 1):
                                lines.append((para.text.strip(), 1, idx))
                        else:
                            continue
                        file_chunks = chunk_text_with_metadata(lines)
                        for chunk in file_chunks:
                            chunks.append((file, chunk["text"], chunk["page"], chunk["line_num"]))
                    except Exception as e:
                        print(f"❌ Failed to process {file_path}: {e}")
        return chunks

    def _get_current_doc_set(self):
        """
        Returns a set of (filename, filesize, mtime) for all supported documents.
        Used to detect changes in the docs folder.
        """
        doc_set = set()
        for root, _, files in os.walk(self.folder_path):
            for file in files:
                ext = os.path.splitext(file)[-1].lower()
                if ext in self.supported_extensions:
                    file_path = os.path.join(root, file)
                    try:
                        stat = os.stat(file_path)
                        doc_set.add((file, stat.st_size, int(stat.st_mtime)))
                    except Exception:
                        continue
        return doc_set

    def _cache_is_stale(self):
        """
        Checks if the cache is stale by comparing doc set with a meta file.
        """
        meta_file = self.cache_file + ".meta"
        current_set = self._get_current_doc_set()
        if not os.path.exists(self.cache_file) or not os.path.exists(meta_file):
            return True
        try:
            with open(meta_file, "r", encoding="utf-8") as f:
                cached_set = set(tuple(x) for x in json.load(f))
            return cached_set != current_set
        except Exception:
            return True

    def _update_meta(self):
        """
        Writes the current doc set to the meta file.
        """
        meta_file = self.cache_file + ".meta"
        doc_set = list(self._get_current_doc_set())
        try:
            with open(meta_file, "w", encoding="utf-8") as f:
                json.dump(doc_set, f)
        except Exception as e:
            print(f"❌ Failed to write chunk cache meta: {e}")

    def refresh_cache(self):
        """
        Force refresh the chunk cache and meta file, syncing with current docs.
        """
        chunks = self._process_and_chunk_documents()
        try:
            with open(self.cache_file, "w", encoding="utf-8") as f:
                for chunk in chunks:
                    f.write(json.dumps(chunk, ensure_ascii=False) + "\n")
            self._update_meta()
            print("✅ Chunk cache refreshed and synced with docs folder.")
        except Exception as e:
            print(f"❌ Failed to refresh chunk cache: {e}")

    def load_documents(self, auto_refresh=True):
        """
        Loads chunks from cache if available and up-to-date, otherwise processes and caches them.
        If auto_refresh is True, will refresh cache if docs folder changes.
        Returns a list of (filename, chunk_text, page, line_num) tuples.
        """
        if auto_refresh and self._cache_is_stale():
            self.refresh_cache()
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, "r", encoding="utf-8") as f:
                    return [tuple(json.loads(line)) for line in f]
            except Exception as e:
                print(f"❌ Failed to load chunk cache: {e}")
        # Fallback: process and cache if cache is missing or failed to load
        self.refresh_cache()
        if os.path.exists(self.cache_file):
            with open(self.cache_file, "r", encoding="utf-8") as f:
                return [tuple(json.loads(line)) for line in f]
        return []
